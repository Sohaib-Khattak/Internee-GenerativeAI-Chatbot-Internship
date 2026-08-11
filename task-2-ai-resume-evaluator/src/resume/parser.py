"""
Resume text extraction from PDF, DOCX, and TXT files.

Extracts clean, readable text from uploaded resume files,
handling various edge cases (password-protected PDFs, scanned PDFs,
large documents, encoding issues).
"""

from __future__ import annotations

import io
import os
import re
from typing import Any, Optional
from werkzeug.datastructures import FileStorage

from config import Config


def extract_text_from_pdf(file_stream: io.BytesIO) -> dict[str, Any]:
    """
    Extract text from a PDF file using PyMuPDF (fitz).

    Args:
        file_stream: BytesIO stream of the PDF file.

    Returns:
        dict with 'text' (str) and 'pages' (int) on success,
        or 'error' (str) and 'code' (str) on failure.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {
            "error": "PDF extraction library not available.",
            "code": "EXTRACTION_FAILED",
        }

    try:
        doc = fitz.open(stream=file_stream, filetype="pdf")

        # Check if PDF is encrypted/password-protected
        if doc.is_encrypted:
            doc.close()
            return {
                "error": "File is password-protected. Please remove protection first.",
                "code": "PROTECTED_FILE",
            }

        total_pages = doc.page_count
        text_parts: list[str] = []

        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            page_text = page.get_text()

            # Detect scanned PDF (very little extractable text)
            if page_num == 0 and len(page_text.strip()) < 20:
                doc.close()
                # Try first few pages to be sure
                full_text = ""
                for p in range(min(total_pages, 3)):
                    full_text += doc.load_page(p).get_text()
                if len(full_text.strip()) < 50:
                    doc.close()
                    return {
                        "error": "No text could be extracted. Try a digital PDF, not a scanned image.",
                        "code": "SCANNED_PDF",
                    }

            text_parts.append(page_text)

        doc.close()
        text = "\n\n".join(text_parts)
        text = normalize_text(text)

        return {"text": text, "pages": total_pages, "format": "pdf"}

    except Exception as exc:
        return {
            "error": f"Couldn't read file: {str(exc)}",
            "code": "EXTRACTION_FAILED",
        }


def extract_text_from_docx(file_stream: io.BytesIO) -> dict[str, Any]:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_stream: BytesIO stream of the DOCX file.

    Returns:
        dict with 'text' (str) and 'pages' (int) on success,
        or 'error' (str) and 'code' (str) on failure.
    """
    try:
        import docx
    except ImportError:
        return {
            "error": "DOCX extraction library not available.",
            "code": "EXTRACTION_FAILED",
        }

    try:
        document = docx.Document(file_stream)
        paragraphs = []

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        text = "\n\n".join(paragraphs)
        text = normalize_text(text)

        if not text.strip():
            return {
                "error": "No text could be extracted from the document.",
                "code": "EXTRACTION_FAILED",
            }

        # Rough page estimate for DOCX (~3000 chars per page)
        page_estimate = max(1, len(text) // 3000)

        return {"text": text, "pages": page_estimate, "format": "docx"}

    except Exception as exc:
        return {
            "error": f"Couldn't read file: {str(exc)}",
            "code": "EXTRACTION_FAILED",
        }


def extract_text_from_txt(file_stream: io.BytesIO) -> dict[str, Any]:
    """
    Extract text from a plain text file.

    Args:
        file_stream: BytesIO stream of the TXT file.

    Returns:
        dict with 'text' (str) and basic metadata.
    """
    try:
        raw = file_stream.read()

        # Try UTF-8 first, fall back to latin-1
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1")

        text = normalize_text(text)

        if not text.strip():
            return {
                "error": "File appears to be empty.",
                "code": "EMPTY_FILE",
            }

        page_estimate = max(1, len(text) // 3000)

        return {"text": text, "pages": page_estimate, "format": "txt"}

    except Exception as exc:
        return {
            "error": f"Couldn't read file: {str(exc)}",
            "code": "EXTRACTION_FAILED",
        }


def extract_text(
    file_storage: FileStorage, filename: str
) -> dict[str, Any]:
    """
    Extract text from an uploaded resume file.

    Dispatches to the appropriate format-specific extractor based on
    file extension. Performs truncation for very long documents.

    Args:
        file_storage: Flask FileStorage object from request.files.
        filename: Sanitized filename with extension.

    Returns:
        dict with keys:
            - 'text' (str): extracted text (on success)
            - 'pages' (int): page count / estimate
            - 'format' (str): file format
            - 'truncated' (bool): whether text was truncated
        OR
            - 'error' (str): user-facing error message
            - 'code' (str): error code for programmatic handling
    """
    ext = os.path.splitext(filename)[1].lower()
    file_stream = io.BytesIO(file_storage.read())

    if ext == ".pdf":
        result = extract_text_from_pdf(file_stream)
    elif ext == ".docx":
        result = extract_text_from_docx(file_stream)
    elif ext == ".txt":
        result = extract_text_from_txt(file_stream)
    else:
        return {
            "error": "Unsupported format. Use PDF, DOCX, or TXT.",
            "code": "UNSUPPORTED_FORMAT",
        }

    if "error" in result:
        return result

    text = result["text"]

    # Truncate very long documents
    truncated = False
    if len(text) > Config.MAX_TEXT_LENGTH:
        text = text[: Config.MAX_TEXT_LENGTH]
        truncated = True

    return {
        "text": text,
        "pages": result.get("pages", 1),
        "format": result.get("format", ext.lstrip(".")),
        "truncated": truncated,
    }


def get_text_preview(text: str, max_chars: int = 500) -> str:
    """
    Get a preview of the extracted text (first N characters).

    Args:
        text: Full extracted text.
        max_chars: Maximum preview length (default: 500).

    Returns:
        Truncated preview string with ellipsis if needed.
    """
    text = text.strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rsplit(" ", 1)[0] + "..."


def normalize_text(text: str) -> str:
    """
    Normalize whitespace in extracted text.

    - Collapses multiple spaces into one
    - Collapses multiple newlines into double newlines (paragraph breaks)
    - Strips leading/trailing whitespace

    Args:
        text: Raw extracted text.

    Returns:
        Normalized text.
    """
    text = re.sub(r" +", " ", text)          # Collapse spaces
    text = re.sub(r"\n{3,}", "\n\n", text)   # Max double newline
    text = re.sub(r"[\r\f\v]", "", text)     # Remove special whitespace
    return text.strip()
